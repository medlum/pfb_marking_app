import zipfile
from pathlib import Path
import shutil
import re
import pandas as pd
import streamlit as st
from docx import Document
from pypdf import PdfReader
import mammoth

def extract_and_read_files(zip_path):
    # Define extraction path
    #extract_folder = "extracted_files"
    extract_folder = st.session_state.user_id
    #st.write("From intern_reflection_report_utils:", st.session_state.user_id)

    if Path(extract_folder).exists():
        shutil.rmtree(extract_folder)

    # Extract ZIP file
    with zipfile.ZipFile(zip_path, "r") as zip_ref:
            zip_ref.extractall(extract_folder)

    extracted_data = {}

    for folder in Path(extract_folder).iterdir(): 

        if folder.is_dir():  

        # extract student name from the each subfolder
        # which contains name as standard label from brightspace 
            folder_name = str(folder.relative_to(extract_folder))
            cleaned_text = re.sub(r"\b(BA|NP|PM|AM)\b", "", folder_name)
            cleaned_text = re.sub(r"\s+", " ", cleaned_text).strip()
            student_name = " ".join(re.findall(r"\b[A-Z]+\b", cleaned_text))

            
            for file in folder.glob("*.*"):

                if file.suffix.lower() == ".docx":
                    #doc = Document(file)
                    #data = "\n".join([para.text for para in doc.paragraphs])
                    #print(data)
                    #data = [[cell.text for cell in row.cells] for table in doc.tables for row in table.rows]

                    def ignore_images(image):
                        return {}
                    
                    result = mammoth.convert_to_html(file,convert_image=ignore_images)
                    data = result.value

                elif file.suffix.lower() == ".pdf":
                    data = ""
                    reader = PdfReader(file)
                    for page in reader.pages:
                        data += page.extract_text()

                else:
                    st.error(".docx or .pdf files not found")
            
                if student_name not in extracted_data:
                    # store values as a list with file extension and extracted data
                    extracted_data[student_name] = [file.suffix.lower(), data]
    
    return extracted_data

#extracted_contents = extract_and_read_files('sample_int6.zip')
#for i in extracted_contents:
#    st.write(extracted_contents[i])

def process_data(data):
    # Convert list of dictionaries to DataFrame
    df = pd.DataFrame(data)

    # Sum the values in 'Program Correctness', 'Code Readability', 'Code Efficiency', 'Documentation', and 'Assignment Specifications'
    df['Total'] = df['Introduction (4 marks)'] + df['OJT Plan (6 marks)'] + df['Analysis and reflection on 3 experiences (Total 30 marks)'] + df[ 'Showcase of accomplished task/achievement (20 marks)'] + df['Diversity and Inclusion (10 marks)'] + df['Influence of internship on future plan (20 marks)'] + df['Quality of writing (10 marks)']
    
    cols = ['Student Name', 
            'Introduction (4 marks)', 
            'OJT Plan (6 marks)', 
            'Analysis and reflection on 3 experiences (Total 30 marks)', 
            'Showcase of accomplished task/achievement (20 marks)', 
            'Diversity and Inclusion (10 marks)',
            'Influence of internship on future plan (20 marks)',
            'Quality of writing (10 marks)',
            'Total', 
            'Feedback']

    return df[cols]

system_message = """
1. Your task is to assess student written assignments using a structured marking rubric.

2. Follow these marking guidelines:
    - Refer closely to the rubric and assign marks per criterion, without exceeding maximum scores.
    - Assign marks with appropriate variation to reflect the quality of each response—avoid giving uniform or overly rounded scores unless well justified.
    - Maintain a high academic standard in grading and feedback.
    - Justify all marks with clear, evidence-based reasoning.

3. Provide detailed, constructive feedback:
    - Structure feedback using line breaks for each major rubric criterion. Use paragraph spacing for clarity.
    - Always refer to “the report” (not “the student”) in your comments.
    - Include at least one direct quote from the report per rubric criterion that requires improvement.
    - For each quote, suggest a revised version that improves formality, clarity, specificity, or conciseness.
    - Follow this format for all sentence-level improvements:
    
        Original: "quoted sentence from the report"  
        Suggestion: "formal, refined version of the sentence"

    - Apply structured frameworks like SMART goals where relevant.
    
    Examples:
        Original: "These are interesting insights and new knowledge to me."  
        Suggestion: "These observations provided valuable insights and deepened my understanding of premium pricing mechanisms."

        Original: "Strengthen my interpersonal skills and expand my network."  
        Suggestion: "By the end of my internship, I aim to initiate at least one meaningful conversation each week with a colleague or mentor to improve my interpersonal communication and expand my professional network."

        Original: "These hands-on experiences have been instrumental in shaping my knowledge and boosting my confidence..."  
        Suggestion: "These hands-on experiences have significantly shaped my knowledge and confidence in client interactions. They also helped me improve how I present myself professionally..."

        Original: "...allowing me to observe how FCs interact with clients..."  
        Suggestion: "...which allowed me to observe how FCs manage client interactions, appointments, and financial products..."

   4. Return your response as a dictionary with the following structure:
        {    
            "Student Name": str,
            "Introduction (4 marks)" : float, 
            "OJT Plan (6 marks)": float, 
            "Analysis and reflection on 3 experiences (Total 30 marks)": float, 
            "Showcase of accomplished task/achievement (20 marks)": float, 
            "Diversity and Inclusion (10 marks)": float,
            "Influence of internship on future plan (20 marks)": float,
            "Quality of writing (10 marks)": float,
            "Feedback": '''Multiline feedback here. Use double quotes inside if quoting student content.'''        
        }
    - Use double quotes (") for all dictionary keys and string values, except for the "Feedback" value.
    - Enclose the "Feedback" value in triple single quotes (''') to preserve formatting and line breaks.
    - Return only the dictionary and nothing else.
"""



# custom CSS for buttons
btn_css = """
<style>
    .stButton > button {
        color: #383736; 
        border: none; /* No border */
        padding: 5px 22px; /* Reduced top and bottom padding */
        text-align: center; /* Centered text */
        text-decoration: none; /* No underline */
        display: inline-block; /* Inline-block */
        font-size: 8px !important;
        margin: 4px 2px; /* Some margin */
        cursor: pointer; /* Pointer cursor on hover */
        border-radius: 30px; /* Rounded corners */
        transition: background-color 0.3s; /* Smooth background transition */
    }
    .stButton > button:hover {
        color: #383736; 
        background-color: #c4c2c0; /* Darker green on hover */
    }
</style>
"""

image_css = """
<style>
.stImage img {
    border-radius: 50%;
    #border: 5px solid #f8fae6;
}
</style>

"""
